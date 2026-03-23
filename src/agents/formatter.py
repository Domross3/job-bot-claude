"""Formatter Agent — renders the approved draft as a polished DOCX resume.

This agent is purely programmatic (no LLM call). It uses python-docx to
generate a professionally styled Word document from the structured
PipelineState data.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..state import PipelineState, ResumeSection

logger = logging.getLogger(__name__)

# ── Style constants ──────────────────────────────────────────────
FONT_NAME = "Calibri"
COLOR_HEADING = RGBColor(0x1A, 0x1A, 0x1A)  # near-black
COLOR_BODY = RGBColor(0x2D, 0x2D, 0x2D)     # dark gray
COLOR_SUBTLE = RGBColor(0x55, 0x55, 0x55)    # medium gray
COLOR_RULE = RGBColor(0x3B, 0x3B, 0x3B)      # dark line


class FormatterAgent:
    """Renders pruned_sections + contact info into a polished .docx file."""

    def run(self, state: PipelineState, output_path: Path) -> PipelineState:
        """Build and save the DOCX resume."""
        logger.info("▶ Running formatter agent [python-docx — no LLM call]")

        doc = Document()
        self._set_page_margins(doc, top=0.5, bottom=0.5, left=0.6, right=0.6)
        self._set_default_font(doc)

        # ── Contact header ───────────────────────────────────────
        name, contact_line = self._extract_contact(state.master_resume)
        self._add_name(doc, name)
        if contact_line:
            self._add_contact_line(doc, contact_line)

        # ── Resume sections ──────────────────────────────────────
        if state.pruned_sections:
            for section in state.pruned_sections:
                self._add_section(doc, section)

        # ── Save ─────────────────────────────────────────────────
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("✔ formatter agent complete — saved to %s", output_path)

        return state.model_copy(update={"final_resume": str(output_path)})

    # ── Contact extraction ───────────────────────────────────────

    @staticmethod
    def _extract_contact(master_resume: str) -> tuple[str, str]:
        """Pull name and contact line from the master resume header."""
        name = "Resume"
        contact_line = ""
        if not master_resume:
            return name, contact_line

        for line in master_resume.split("\n"):
            stripped = line.strip()
            if stripped.startswith("---"):
                break
            # First heading = name
            if stripped.startswith("#") and name == "Resume":
                name = stripped.lstrip("#").strip()
                # Remove markdown formatting artifacts
                name = re.sub(r"[*_]", "", name).strip()
                # Strip common suffixes like "— Master Resume", "- Resume"
                name = re.sub(r"\s*[—–-]\s*(Master\s+)?Resume.*$", "", name, flags=re.IGNORECASE).strip()
            # Line with | or @ = contact info
            elif ("|" in stripped or "@" in stripped) and name != "Resume":
                # Strip markdown bold/italic markers and link syntax
                contact_line = re.sub(r"\*\*|\*|__|_", "", stripped)
                contact_line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", contact_line)
                contact_line = contact_line.strip()

        return name, contact_line

    # ── Document builders ────────────────────────────────────────

    def _add_name(self, doc: Document, name: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(0)
        p.space_after = Pt(2)
        run = p.add_run(name)
        run.font.name = FONT_NAME
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = COLOR_HEADING

    def _add_contact_line(self, doc: Document, contact: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(0)
        p.space_after = Pt(6)
        run = p.add_run(contact)
        run.font.name = FONT_NAME
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_SUBTLE

    def _add_section(self, doc: Document, section: ResumeSection) -> None:
        """Add a full resume section (heading + entries)."""
        heading = section.heading.upper()

        # ── Section heading with bottom rule ─────────────────────
        p = doc.add_paragraph()
        p.space_before = Pt(10)
        p.space_after = Pt(3)
        run = p.add_run(heading)
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_HEADING
        # Add a thin bottom border
        self._add_bottom_border(p)

        # ── Entries ──────────────────────────────────────────────
        for entry in section.entries:
            self._add_entry(doc, entry, section.heading)

    def _add_entry(self, doc: Document, entry, section_heading: str) -> None:
        """Add a single entry (role, degree, project, etc.)."""
        is_skills = "skill" in section_heading.lower()

        if is_skills:
            # Skills: render bullets as styled lines (no title/org/dates)
            for bullet in entry.bullets:
                p = doc.add_paragraph()
                p.space_before = Pt(1)
                p.space_after = Pt(1)
                p.paragraph_format.left_indent = Inches(0.15)

                # If bullet has "Category: items" format, bold the category
                if ":" in bullet:
                    category, items = bullet.split(":", 1)
                    cat_run = p.add_run(f"{category.strip()}:")
                    cat_run.font.name = FONT_NAME
                    cat_run.font.size = Pt(10)
                    cat_run.font.bold = True
                    cat_run.font.color.rgb = COLOR_BODY
                    items_run = p.add_run(f" {items.strip()}")
                    items_run.font.name = FONT_NAME
                    items_run.font.size = Pt(10)
                    items_run.font.color.rgb = COLOR_BODY
                else:
                    run = p.add_run(bullet)
                    run.font.name = FONT_NAME
                    run.font.size = Pt(10)
                    run.font.color.rgb = COLOR_BODY
            return

        # ── Title | Organization | Dates line ────────────────────
        p = doc.add_paragraph()
        p.space_before = Pt(5)
        p.space_after = Pt(1)

        # Bold title
        title_run = p.add_run(entry.title)
        title_run.font.name = FONT_NAME
        title_run.font.size = Pt(10.5)
        title_run.font.bold = True
        title_run.font.color.rgb = COLOR_HEADING

        # Separator + organization
        if entry.organization:
            sep_run = p.add_run("  |  ")
            sep_run.font.name = FONT_NAME
            sep_run.font.size = Pt(10.5)
            sep_run.font.color.rgb = COLOR_SUBTLE

            org_run = p.add_run(entry.organization)
            org_run.font.name = FONT_NAME
            org_run.font.size = Pt(10.5)
            org_run.font.color.rgb = COLOR_BODY

        # Dates (right-aligned feel via tab — or appended)
        if entry.dates:
            sep_run = p.add_run("  |  ")
            sep_run.font.name = FONT_NAME
            sep_run.font.size = Pt(10.5)
            sep_run.font.color.rgb = COLOR_SUBTLE

            date_run = p.add_run(entry.dates)
            date_run.font.name = FONT_NAME
            date_run.font.size = Pt(10.5)
            date_run.font.italic = True
            date_run.font.color.rgb = COLOR_SUBTLE

        # ── Bullet points ────────────────────────────────────────
        for bullet in entry.bullets:
            bp = doc.add_paragraph()
            bp.space_before = Pt(0.5)
            bp.space_after = Pt(0.5)
            bp.paragraph_format.left_indent = Inches(0.25)
            bp.paragraph_format.first_line_indent = Inches(-0.15)

            # Bullet character
            marker_run = bp.add_run("•  ")
            marker_run.font.name = FONT_NAME
            marker_run.font.size = Pt(10)
            marker_run.font.color.rgb = COLOR_SUBTLE

            text_run = bp.add_run(bullet)
            text_run.font.name = FONT_NAME
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = COLOR_BODY

    # ── Low-level helpers ────────────────────────────────────────

    @staticmethod
    def _set_page_margins(
        doc: Document, top: float, bottom: float, left: float, right: float
    ) -> None:
        """Set page margins in inches."""
        for section in doc.sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)

    @staticmethod
    def _set_default_font(doc: Document) -> None:
        """Set the document-wide default font."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = FONT_NAME
        font.size = Pt(10)
        font.color.rgb = COLOR_BODY
        # Tight line spacing
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = Pt(13)

    @staticmethod
    def _add_bottom_border(paragraph) -> None:
        """Add a thin bottom border to a paragraph (section heading rule)."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(
            qn("w:bottom"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "4",        # 0.5pt line
                qn("w:space"): "1",
                qn("w:color"): "3B3B3B",
            },
        )
        pBdr.append(bottom)
        pPr.append(pBdr)
